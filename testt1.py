from docx import Document
import openpyxl

from func_genitive import genitive, genitive_name, number_to_words

theFile = openpyxl.load_workbook('Сертификация исходник.xlsx')
currentSheet = theFile['Сертификация']

MONTHS = {
    '01': 'январь',
    '02': 'февраль',
    '03': 'март',
    '04': 'апрель',
    '05': 'май',
    '06': 'июнь',
    '07': 'июль',
    '08': 'август',
    '09': 'сентябрь',
    '10': 'октябрь',
    '11': 'ноябрь',
    '12': 'декабрь',
}


variables = {
    '{{CONTRACT_NUMBER}}': str(currentSheet['B2'].value),
    '{{CONTRACT_YEAR}}': str(currentSheet['B3'].value),  # Короткий год - 24
    '{{BILL_NUMBER}}': str(currentSheet['B4'].value),  # Короткий год - 24
    '{{CONTR_DATE}}': str(currentSheet['B5'].value),
    '{{CONTR_MONTH}}': str(currentSheet['B6'].value),
    #  Если месяц договора декабрь, то месяц окончания работ - январь, чтобы не получилось месяц окончания - 13
    '{{CONTR_MONTH+1}}': '01' if str(currentSheet['B6'].value) == '12' else str(int(currentSheet['B6'].value) + 1),  # Для указания месяца проведения работ.
    '{{CONTR_MONTH_WORD_GEN}}': genitive(MONTHS[str(currentSheet['B6'].value)]),
    '{{CONTR_YEAR}}': '20' + str(currentSheet['B3'].value),  # Длинный год - 2024
    #  Равен году договора, а если месяц договора декабрь, то год договора + 1
    '{{CONTR_YEAR+1}}': '20' + str(currentSheet['B3'].value) if str(currentSheet['B6'].value) != '12' else '20' + str(int(currentSheet['B3'].value) + 1),  # Для указания года окончания проведения работ.
    '{{BUSINESS_FORM_FULL}}': str(currentSheet['B7'].value),
    '{{COMPANY_NAME_FULL}}': str(currentSheet['B8'].value),
    '{{BUSINESS_FORM}}': str(currentSheet['B9'].value),
    '{{COMPANY_NAME}}': str(currentSheet['B10'].value),
    '{{DIR_LASTNAME}}': str(currentSheet['B11'].value),  # Фамилия
    '{{DIR_FIRSTNAME}}': str(currentSheet['B12'].value),  # Имя
    '{{DIR_SECNAME}}': str(currentSheet['B13'].value),  # Отчество
    '{{DIR_FIRSTNAME_SHORT}}': str(currentSheet['B12'].value)[:1] + '.',
    '{{DIR_SECNAME_SHORT}}': str(currentSheet['B13'].value)[:1] + '.',
    '{{DIR_LASTNAME_GEN}}': genitive_name('LASTNAME', str(currentSheet['B14'].value), str(currentSheet['B11'].value)),  # Фамилия
    '{{DIR_FIRSTNAME_GEN}}': genitive_name('FIRSTNAME', str(currentSheet['B14'].value), str(currentSheet['B12'].value)),  # Имя
    '{{DIR_SECNAME_GEN}}': genitive_name('MIDDLENAME', str(currentSheet['B14'].value), str(currentSheet['B13'].value)),  # Отчество
    '{{CERT_NAME}}': str(currentSheet['B15'].value),
    '{{CERT_GROUP}}': str(currentSheet['B16'].value),
    '{{OKPD}}': str(currentSheet['B17'].value),
    '{{STANDART_MAIN}}': str(currentSheet['B18'].value),
    '{{STANDART_SHORT}}': str(currentSheet['B19'].value),
    '{{STANDART_FULL}}': str(currentSheet['B20'].value),
    '{{CONTRACT_SUM}}': str(currentSheet['B21'].value),
    '{{CONTRACT_OS_FULL_SUM}}': str(currentSheet['B22'].value),
    '{{CONTRACT_OS_SUM}}': str(currentSheet['B23'].value),
    '{{CONTRACT_IL_FULL_SUM}}': str(currentSheet['B24'].value),
    '{{CONTRACT_IL_SUM}}': str(currentSheet['B25'].value),
    '{{CONTRACT_IK_SUM}}': str(currentSheet['B26'].value),
    '{{CONTRACT_SUM_WORDS}}': number_to_words(str(currentSheet['B21'].value)),
    '{{INN}}': str(currentSheet['B27'].value),
    '{{KPP}}': str(currentSheet['B28'].value),
    '{{JUR_ADDRESS}}': str(currentSheet['B29'].value),
    '{{PHYS_ADDRESS}}': str(currentSheet['B30'].value),
    '{{RAS_ACCOUNT}}': str(currentSheet['B31'].value),
    '{{BANK_NAME}}': str(currentSheet['B32'].value),
    '{{KORR_ACCOUNT}}': str(currentSheet['B33'].value),
    '{{BIK}}': str(currentSheet['B34'].value),
    '{{OGRN}}': str(currentSheet['B35'].value),
    '{{TEL}}': str(currentSheet['B36'].value),
    '{{E-MAIL}}': str(currentSheet['B37'].value),
    '{{EXPERT_LASTNAME}}': str(currentSheet['B38'].value),  # Фамилия
    '{{EXPERT_FIRSTNAME_SHORT}}': str(currentSheet['B39'].value),
    '{{EXPERT_SECNAME_SHORT}}': str(currentSheet['B40'].value),
    '{{EXPERT_REG_NUMBER}}': str(currentSheet['B41'].value),
    '{{IL_NAME}}': str(currentSheet['B42'].value),
    '{{IL_REG_NUMBER}}': str(currentSheet['B43'].value),
    '{{IL_EXPIRE_DATE}}': str(currentSheet['B44'].value),
    '{{ISSUE_DECISION_DATE}}': str(currentSheet['B45'].value),
    '{{ISSUE_DECISION_MONTH}}': str(currentSheet['B46'].value),
    '{{ISSUE_DECISION_YEAR}}': str(currentSheet['B47'].value),
    '{{CERTIFICARTE_START_DATE}}': str(currentSheet['B48'].value),
    '{{CERTIFICARTE_START_MONTH}}': str(currentSheet['B49'].value),
    '{{CERTIFICARTE_START_YEAR}}': str(currentSheet['B50'].value),
    '{{CERTIFICARTE_DURATION}}': str(currentSheet['B51'].value),
    '{{SAMPLE_ACT_DATE}}': str(currentSheet['B52'].value),
    '{{SAMPLE_ACT_MONTH}}': str(currentSheet['B53'].value),
    '{{SAMPLE_ACT_MONTH_GEN}}': genitive(MONTHS[str(currentSheet['B53'].value)]),
    '{{SAMPLE_ACT_YEAR}}': str(currentSheet['B54'].value),
    '{{PRODUCTION_CREATED_QUARTER}}': str(currentSheet['B55'].value),
    '{{PRODUCTION_CREATED_YEAR}}': str(currentSheet['B56'].value),
    '{{PRODUCTION_SAMPLES}}': str(currentSheet['B57'].value),
    '{{PROTOCOL_DATE}}': str(currentSheet['B58'].value),
    '{{PROTOCOL_MONTH}}': str(currentSheet['B59'].value),
    '{{PROTOCOL_MONTH_WORD_GEN}}': genitive(MONTHS[str(currentSheet['B59'].value)]),
    '{{PROTOCOL_YEAR}}': str(currentSheet['B60'].value),
    '{{PROTOCOL_NUMBER}}': str(currentSheet['B61'].value),
    '{{TEST_GOSTS}}': str(currentSheet['B62'].value),
    '{{TEST_START_DATE}}': str(currentSheet['B63'].value),
    '{{TEST_START_MONTH}}': str(currentSheet['B64'].value),
    '{{TEST_START_YEAR}}': str(currentSheet['B65'].value),
    '{{SAMPLES_MARK}}': str(currentSheet['B66'].value),
    '{{TESTER_NAME}}': str(currentSheet['B67'].value),
    '{{PROD_ANALYSE_DATE}}': str(currentSheet['B68'].value),
    '{{PROD_ANALYSE_MONTH}}': str(currentSheet['B69'].value),
    '{{PROD_ANALYSE_YEAR}}': str(currentSheet['B70'].value),
}


def main(template_file_path, output_file_path, variables):
    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    print('OK!!!!!!!!!!!!!!!!!!!', template_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        text_runs = [(run.text, run.font.name, run.font.size, run.bold, run.italic, run.underline) for run in paragraph.runs]
        full_text = ''.join([text for text, _, _, _, _, _ in text_runs])
        new_text = full_text.replace(key, value)
        
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''

        # Create a new run to hold the entire new text
        new_run = paragraph.add_run(new_text)
        
        # Apply formatting from the first run
        if text_runs:
            _, font_name, font_size, bold, italic, underline = text_runs[0]
            new_run.font.name = font_name or 'Times New Roman'
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline


def reference_data_xlsx(input_reference, output_reference):
    theFile = openpyxl.load_workbook(input_reference)
    currentSheet = theFile['На регистрацию сертиф.']
    currentSheet['D8'].value = f"{variables['{{CONTRACT_NUMBER}}']} от {variables['{{CONTR_DATE}}']}.{variables['{{CONTR_MONTH}}']}.{variables['{{CONTR_YEAR}}']}г."
    currentSheet['D11'].value = f"{variables['{{EXPERT_LASTNAME}}']} {variables['{{EXPERT_FIRSTNAME_SHORT}}']}{variables['{{EXPERT_SECNAME_SHORT}}']}"
    currentSheet['D12'].value = variables['{{EXPERT_REG_NUMBER}}']
    currentSheet['D13'].value = variables['{{IL_NAME}}']
    currentSheet['D14'].value = variables['{{IL_REG_NUMBER}}']
    currentSheet['D15'].value = variables['{{BUSINESS_FORM}}']
    currentSheet['D16'].value = variables['{{COMPANY_NAME}}']
    currentSheet['D17'].value = variables['{{JUR_ADDRESS}}']
    currentSheet['D18'].value = variables['{{TEL}}']
    currentSheet['D19'].value = variables['{{E-MAIL}}']
    currentSheet['D20'].value = f"{variables['{{DIR_LASTNAME}}']} {variables['{{DIR_FIRSTNAME}}']} {variables['{{DIR_SECNAME}}']}"
    currentSheet['D21'].value = variables['{{INN}}']
    currentSheet['D22'].value = variables['{{KPP}}']
    currentSheet['D26'].value = variables['{{CERT_GROUP}}']
    currentSheet['D27'].value = variables['{{CERT_NAME}}']
    currentSheet['D28'].value = variables['{{OKPD}}']
    currentSheet['D31'].value = variables['{{STANDART_SHORT}}']
    currentSheet['D32'].value = variables['{{STANDART_FULL}}']
    currentSheet['D33'].value = variables['{{CONTRACT_NUMBER}}']
    currentSheet['D34'].value = f"{variables['{{ISSUE_DECISION_DATE}}']}.{variables['{{ISSUE_DECISION_MONTH}}']}.{variables['{{ISSUE_DECISION_YEAR}}']}"
    currentSheet['D35'].value = f"{variables['{{CERTIFICARTE_START_DATE}}']}.{variables['{{CERTIFICARTE_START_MONTH}}']}.{variables['{{CERTIFICARTE_START_YEAR}}']}"
    currentSheet['D36'].value = f"{variables['{{CERTIFICARTE_DURATION}}']} {'года' if int(variables['{{CERTIFICARTE_DURATION}}']) <= 4 else 'лет'}"
    currentSheet['D37'].value = f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 1)}"
    currentSheet['D38'].value = f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 2)}"
    currentSheet['D39'].value = "" if int(variables['{{CERTIFICARTE_DURATION}}']) <= 3 else f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 3)}"
    currentSheet['D40'].value = "" if int(variables['{{CERTIFICARTE_DURATION}}']) <= 4 else f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 4)}"
    currentSheet['D41'].value = f"{variables['{{SAMPLE_ACT_DATE}}']}.{variables['{{SAMPLE_ACT_MONTH}}']}.{variables['{{SAMPLE_ACT_YEAR}}']}"
    currentSheet['D43'].value = variables['{{PROTOCOL_NUMBER}}']
    currentSheet['D44'].value = f"{variables['{{PROTOCOL_DATE}}']}.{variables['{{PROTOCOL_MONTH}}']}.{variables['{{PROTOCOL_YEAR}}']}"
    currentSheet['D45'].value = f"{variables['{{TESTER_NAME}}']}"
    currentSheet['D46'].value = f"{variables['{{PROD_ANALYSE_DATE}}']}.{variables['{{PROD_ANALYSE_MONTH}}']}.{variables['{{PROD_ANALYSE_YEAR}}']} № {variables['{{CONTRACT_NUMBER}}']}"
    currentSheet['D47'].value = f"{variables['{{CONTRACT_NUMBER}}'][:2]}-{variables['{{CONTRACT_YEAR}}']}"
    currentSheet['D48'].value = f"{variables['{{CONTR_DATE}}']}.{variables['{{CONTR_MONTH}}']}.{variables['{{CONTR_YEAR}}']}г."
    currentSheet['D49'].value = f"{variables['{{CONTRACT_SUM}}']}"
    currentSheet['D50'].value = f"{variables['{{CONTRACT_OS_FULL_SUM}}']}"
    currentSheet['D51'].value = f"{variables['{{CONTRACT_OS_SUM}}']}"
    currentSheet['D52'].value = f"{variables['{{CONTRACT_IL_FULL_SUM}}']}"
    currentSheet['D53'].value = f"{variables['{{CONTRACT_IL_SUM}}']}"
    theFile.save(output_reference)

    theFile_Insp = openpyxl.load_workbook('Инспекция.xlsx')
    currentSheet = theFile_Insp['Инспекция']
    currentSheet['D13'].value = 123
    theFile_Insp.save("Всемаф Инспекция.xlsx")



if __name__ == '__main__':
    #  main('00 Подписанный Дог серт Исходник.docx', f"00 Подписанный Дог серт {str(currentSheet['B10'].value)}.docx", variables)
    #  main('01 Дог серт Исходник.docx', f"01 Дог серт {str(currentSheet['B10'].value)}.docx", variables)
    #  main('02 Счет Исходник.docx', f"02 Счет {str(currentSheet['B10'].value)}.docx", variables)
    # main('03 Договор ИК Исходник.docx', f"03 Договор ИК {str(currentSheet['B10'].value)}.docx", variables)
    # main('04 АКТ приемки Исходник.docx', f"04 АКТ приемки {str(currentSheet['B10'].value)}.docx", variables)
    main('01 Заявка Исходник.docx', f"01 Заявка {str(currentSheet['B10'].value)}.docx", variables)
    main('02 Распоряжение Исходник.docx', f"02 Распоряжение {str(currentSheet['B10'].value)}.docx", variables)
    main('03 Акт отбора Исходник.docx', f"03 Акт отбора {str(currentSheet['B10'].value)}.docx", variables)
    main('03 Протокол Исходник.docx', f"03 Протокол {str(currentSheet['B10'].value)}.docx", variables)
    main('04 Анализ производства Исходник.docx', f"04 Анализ производства {str(currentSheet['B10'].value)}.docx", variables)
    main('07 Макет сертификата Исходник.docx', f"07 Макет сертификата {str(currentSheet['B10'].value)}.docx", variables)
    reference_data_xlsx('08 Справка Исходник.xlsx', f"08 Справка {str(currentSheet['B10'].value)}.xlsx")
