import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from docx import Document
from certification_many_gui import main_func

# Function to replace text in a paragraph
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        text_runs = [(run.text, run.font.name, run.font.size, run.bold, run.italic, run.underline) for run in paragraph.runs]
        full_text = ''.join([text for text, _, _, _, _, _ in text_runs])
        new_text = full_text.replace(key, value)

        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''

        # Create a new run to hold the entire new text
        new_run = paragraph.add_run(new_text)

        # Apply formatting from the first run
        if text_runs:
            _, font_name, font_size, bold, italic, underline = text_runs[0]
            new_run.font.name = font_name or 'Times New Roman'
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline

# Function to process the document with the given template and variables
def main(template_file_path, output_file_path, variables):
    template_document = Document(template_file_path)

    for variable_set in variables.values():
        for variable_key, variable_value in variable_set.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    messagebox.showinfo("Success", f"Document saved as {output_file_path}")

# Function to get file paths from the user
def get_file_paths():
    template_file_path = filedialog.askopenfilename(title="Select Template File", filetypes=[("Word Documents", "*.docx")])
    output_file_path = filedialog.asksaveasfilename(title="Save Output File", defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    return template_file_path, output_file_path

# Function to handle the OK button click
def on_ok():
    variables = {}
    for i, tab in enumerate(tabs):
        contract_number = tab['CONTRACT_NUMBER'].get()
        contract_year = tab['CONTRACT_YEAR'].get()
        work_number = tab['WORK_NUMBER'].get()
        
        variables[i] = {
            '{{CONTRACT_NUMBER}}': contract_number,
            '{{CONTRACT_YEAR}}': contract_year,
            '{{WORK_NUMBER}}': work_number
        }

    main_func(variables)

# Create the main window
root = tk.Tk()
root.title("Document Generator")

# Create a style for the "+" button
style = ttk.Style()
style.configure('Plus.TButton', font=('Arial', 12, 'bold'), foreground='blue')

# Create a canvas to hold the notebook and the "+" button
canvas = tk.Canvas(root)
canvas.pack(expand=True, fill='both')

# Create a Notebook widget
notebook = ttk.Notebook(canvas)
notebook.pack(side='left', expand=True, fill='both')

# List to store tab data
tabs = []

# Function to add a new tab
def add_tab():
    tab_frame = ttk.Frame(notebook)
    tab_data = {}

    tk.Label(tab_frame, text="CONTRACT_NUMBER:").grid(row=0, column=0, padx=10, pady=5)
    contract_number_entry = tk.Entry(tab_frame)
    contract_number_entry.grid(row=0, column=1, padx=10, pady=5)
    tab_data['CONTRACT_NUMBER'] = contract_number_entry

    tk.Label(tab_frame, text="CONTRACT_YEAR:").grid(row=1, column=0, padx=10, pady=5)
    contract_year_entry = tk.Entry(tab_frame)
    contract_year_entry.grid(row=1, column=1, padx=10, pady=5)
    tab_data['CONTRACT_YEAR'] = contract_year_entry

    tk.Label(tab_frame, text="WORK_NUMBER:").grid(row=2, column=0, padx=10, pady=5)
    work_number_entry = tk.Entry(tab_frame)
    work_number_entry.grid(row=2, column=1, padx=10, pady=5)
    tab_data['WORK_NUMBER'] = work_number_entry

    tabs.append(tab_data)
    notebook.add(tab_frame, text=f'Tab {len(tabs)}')
    update_plus_button_position()

# Function to update the "+" button position
def update_plus_button_position():
    canvas.update_idletasks()
    x = notebook.winfo_width()
    canvas.coords(plus_button_id, x + 5, 5)

# Create and place the "+" button
plus_button = ttk.Button(canvas, text="+", style='Plus.TButton', command=add_tab)
plus_button_id = canvas.create_window(10, 10, window=plus_button, anchor='nw')

# Add initial tabs
add_tab()
add_tab()

# Pack the canvas and place the "+" button in the correct position
canvas.pack(expand=True, fill='both')
root.update_idletasks()
update_plus_button_position()

# Create and place the OK button
ok_button = tk.Button(root, text="OK", command=on_ok)
ok_button.pack(pady=10)

# Run the main event loop
root.mainloop()
