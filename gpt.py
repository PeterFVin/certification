from docx import Document
import openpyxl

theFile = openpyxl.load_workbook('excel_file.xlsx')
currentSheet = theFile['excel_sheet']

variables = {
    '{{CONTRACT_NUMBER}}': str(currentSheet['B2'].value),
    '{{CONTRACT_YEAR}}': str(currentSheet['B3'].value),
    '{{WORK_NUMBER}}': str(currentSheet['B4'].value),
    '{{BILL_NUMBER}}': str(currentSheet['B5'].value),
    '{{CONTR_DATE}}': str(currentSheet['B6'].value),
    '{{CONTR_MONTH}}': str(currentSheet['B7'].value),
}


def main(template_file_path, output_file_path, variables):
    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    print('OK!!!!!!!!!!!!!!!!!!!', template_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        text_runs = [(run.text, run.font.name, run.font.size, run.bold, run.italic, run.underline) for run in paragraph.runs]
        full_text = ''.join([text for text, _, _, _, _, _ in text_runs])
        new_text = full_text.replace(key, value)
        
        for run in paragraph.runs:
            run.text = ''

        new_run = paragraph.add_run(new_text)
        
        if text_runs:
            _, font_name, font_size, bold, italic, underline = text_runs[0]
            new_run.font.name = font_name or 'Times New Roman'
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline


if __name__ == '__main__':
    main('input_file.docx', f"output_file.docx", variables)



currentSheet = theFile[excel_sheet_change]

