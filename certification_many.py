from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import openpyxl

from func_genitive import genitive, genitive_name, number_to_words


MONTHS = {
    '01': 'январь',
    '02': 'февраль',
    '03': 'март',
    '04': 'апрель',
    '05': 'май',
    '06': 'июнь',
    '07': 'июль',
    '08': 'август',
    '09': 'сентябрь',
    '10': 'октябрь',
    '11': 'ноябрь',
    '12': 'декабрь',
}

theFile = openpyxl.load_workbook('Сертификация много исходник.xlsx')


def load_variables_from_sheet(excel_file_path, sheet_number):
    currentSheet = theFile.worksheets[sheet_number]

    # variables1 = {}

    # try:
    #     value = int(currentSheet['B2'].value)
    # except Exception:
    #     raise Exception('Ошибка в ячейке B2: введите целое положительное число!')
    # variables1['{{CONTRACT_NUMBER}}'] = str(currentSheet['B2'].value)
    # print(variables1)

    variables = {
        '{{CONTRACT_NUMBER}}': str(currentSheet['B2'].value),
        '{{CONTRACT_YEAR}}': str(currentSheet['B3'].value),  # Короткий год - 24
        '{{WORK_NUMBER}}': str(currentSheet['B4'].value),
        '{{BILL_NUMBER}}': str(currentSheet['B5'].value),  # Короткий год - 24
        '{{CONTR_DATE}}': str(currentSheet['B6'].value),
        #  Для распоряжения. Ограничение максимального дня- 28, чтобы не получилось 29.02
        '{{CONTR_DATE<29}}': str(currentSheet['B6'].value) if int(currentSheet['B6'].value) < 29 else '28',
        '{{CONTR_MONTH}}': str(currentSheet['B7'].value),
        #  Для распоряжения. Если месяц договора декабрь, то месяц окончания работ - январь, чтобы не получилось месяц окончания - 13
        '{{CONTR_MONTH+1}}': '01' if str(currentSheet['B7'].value) == '12' else str(int(currentSheet['B7'].value) + 1),  # Для указания месяца проведения работ.
        '{{CONTR_MONTH_WORD_GEN}}': genitive(MONTHS[str(currentSheet['B7'].value)]),
        '{{CONTR_YEAR}}': '20' + str(currentSheet['B3'].value),  # Длинный год - 2024
        #  Равен году договора, а если месяц договора декабрь, то год договора + 1
        '{{CONTR_YEAR+1}}': '20' + str(currentSheet['B3'].value) if str(currentSheet['B7'].value) != '12' else '20' + str(int(currentSheet['B3'].value) + 1),  # Для указания года окончания проведения работ.
        '{{BUSINESS_FORM_FULL}}': str(currentSheet['B8'].value),
        '{{COMPANY_NAME_FULL}}': str(currentSheet['B9'].value),
        '{{BUSINESS_FORM}}': str(currentSheet['B10'].value),
        '{{COMPANY_NAME}}': str(currentSheet['B11'].value),
        '{{DIR_LASTNAME}}': str(currentSheet['B12'].value),  # Фамилия
        '{{DIR_FIRSTNAME}}': str(currentSheet['B13'].value),  # Имя
        '{{DIR_SECNAME}}': str(currentSheet['B14'].value),  # Отчество
        '{{DIR_FIRSTNAME_SHORT}}': str(currentSheet['B13'].value)[:1] + '.',
        '{{DIR_SECNAME_SHORT}}': str(currentSheet['B14'].value)[:1] + '.',
        '{{DIR_LASTNAME_GEN}}': genitive_name('LASTNAME', str(currentSheet['B15'].value), str(currentSheet['B12'].value)),  # Фамилия
        '{{DIR_FIRSTNAME_GEN}}': genitive_name('FIRSTNAME', str(currentSheet['B15'].value), str(currentSheet['B13'].value)),  # Имя
        '{{DIR_SECNAME_GEN}}': genitive_name('MIDDLENAME', str(currentSheet['B15'].value), str(currentSheet['B14'].value)),  # Отчество
        '{{CERT_NAME}}': str(currentSheet['B16'].value),
        '{{CERT_GROUP}}': str(currentSheet['B17'].value),
        '{{OKPD}}': str(currentSheet['B18'].value),
        '{{STANDART_MAIN}}': str(currentSheet['B19'].value),
        '{{STANDART_SHORT}}': str(currentSheet['B20'].value),
        '{{STANDART_FULL}}': str(currentSheet['B21'].value),
        '{{CONTRACT_SUM}}': str(currentSheet['B22'].value),
        '{{CONTRACT_OS_FULL_SUM}}': str(currentSheet['B23'].value),
        '{{CONTRACT_OS_SUM}}': str(currentSheet['B24'].value),
        '{{CONTRACT_IL_FULL_SUM}}': str(currentSheet['B25'].value),
        '{{CONTRACT_IL_SUM}}': str(currentSheet['B26'].value),
        '{{CONTRACT_IK_SUM}}': str(currentSheet['B27'].value),
        '{{CONTRACT_SUM_WORDS}}': number_to_words(str(currentSheet['B22'].value)),
        '{{INN}}': str(currentSheet['B28'].value),
        '{{KPP}}': str(currentSheet['B29'].value),
        '{{JUR_ADDRESS}}': str(currentSheet['B30'].value),
        '{{PHYS_ADDRESS}}': str(currentSheet['B31'].value),
        '{{RAS_ACCOUNT}}': str(currentSheet['B32'].value),
        '{{BANK_NAME}}': str(currentSheet['B33'].value),
        '{{KORR_ACCOUNT}}': str(currentSheet['B34'].value),
        '{{BIK}}': str(currentSheet['B35'].value),
        '{{OGRN}}': str(currentSheet['B36'].value),
        '{{TEL}}': str(currentSheet['B37'].value),
        '{{E-MAIL}}': str(currentSheet['B38'].value),
        '{{EXPERT_LASTNAME}}': str(currentSheet['B39'].value),  # Фамилия
        '{{EXPERT_FIRSTNAME_SHORT}}': str(currentSheet['B40'].value),
        '{{EXPERT_SECNAME_SHORT}}': str(currentSheet['B41'].value),
        '{{EXPERT_REG_NUMBER}}': str(currentSheet['B42'].value),
        '{{IL_NAME}}': str(currentSheet['B43'].value),
        '{{IL_REG_NUMBER}}': str(currentSheet['B44'].value),
        '{{IL_EXPIRE_DATE}}': str(currentSheet['B45'].value),
        '{{ISSUE_DECISION_DATE}}': str(currentSheet['B46'].value),
        '{{ISSUE_DECISION_MONTH}}': str(currentSheet['B47'].value),
        '{{ISSUE_DECISION_MONTH_GEN}}': genitive(MONTHS[str(currentSheet['B47'].value)]),
        '{{ISSUE_DECISION_YEAR}}': str(currentSheet['B48'].value),
        '{{CERTIFICARTE_START_DATE}}': str(currentSheet['B49'].value),
        '{{CERTIFICARTE_START_MONTH}}': str(currentSheet['B50'].value),
        '{{CERTIFICARTE_START_YEAR}}': str(currentSheet['B51'].value),
        '{{CERTIFICARTE_DURATION}}': str(currentSheet['B52'].value),
        '{{SAMPLE_ACT_DATE}}': str(currentSheet['B53'].value),
        '{{SAMPLE_ACT_MONTH}}': str(currentSheet['B54'].value),
        '{{SAMPLE_ACT_MONTH_GEN}}': genitive(MONTHS[str(currentSheet['B54'].value)]),
        '{{SAMPLE_ACT_YEAR}}': str(currentSheet['B55'].value),
        '{{PRODUCTION_CREATED_QUARTER}}': str(currentSheet['B56'].value),
        '{{PRODUCTION_CREATED_YEAR}}': str(currentSheet['B57'].value),
        '{{PRODUCTION_SAMPLES}}': str(currentSheet['B58'].value),
        '{{PROTOCOL_DATE}}': str(currentSheet['B59'].value),
        '{{PROTOCOL_MONTH}}': str(currentSheet['B60'].value),
        '{{PROTOCOL_MONTH_WORD_GEN}}': genitive(MONTHS[str(currentSheet['B60'].value)]),
        '{{PROTOCOL_YEAR}}': str(currentSheet['B61'].value),
        '{{PROTOCOL_NUMBER}}': str(currentSheet['B62'].value),
        '{{TEST_GOSTS}}': str(currentSheet['B63'].value),
        '{{TEST_START_DATE}}': str(currentSheet['B64'].value),
        '{{TEST_START_MONTH}}': str(currentSheet['B65'].value),
        '{{TEST_START_YEAR}}': str(currentSheet['B66'].value),
        '{{SAMPLES_MARK}}': str(currentSheet['B67'].value),
        '{{TESTER_NAME}}': str(currentSheet['B68'].value),
        '{{PROD_ANALYSE_DATE}}': str(currentSheet['B69'].value),
        '{{PROD_ANALYSE_MONTH}}': str(currentSheet['B70'].value),
        '{{PROD_ANALYSE_YEAR}}': str(currentSheet['B71'].value),
        '{{IK&CERT_DAY}}': str(currentSheet['B72'].value),
        '{{IK&CERT_MONTH}}': str(currentSheet['B73'].value),
        '{{IK&CERT_WORD_MONTH}}': genitive(MONTHS[str(currentSheet['B73'].value)]),
        '{{IK&CERT_YEAR}}': str(currentSheet['B74'].value),
        'worksheets_count': len(theFile.worksheets),
    }

    value = str(currentSheet['B27'].value)
    match int(currentSheet['B52'].value):
        case 4:
            variables['{{IK_ADD_SUM}}'] = f"\n\tIII этап\t\t{value} руб. 00 коп."
        case 5:
            variables['{{IK_ADD_SUM}}'] = f"\n\tIII этап\t\t{value} руб. 00 коп.\n\tVI этап\t\t{value} руб. 00 коп."
        case _:
            variables['{{IK_ADD_SUM}}'] = ''
    
    # variables['cert_list'] = []
    # for work in range(variables['worksheets_count']):
    #     variables['cert_list'].append()

    return variables


def main(template_file_path, output_file_path, variables):
    template_document = Document(template_file_path)

    # print(variables)
    print('!!!!!!!!!!!!!!!!!!!!!!!!!!')
    # print(variables.items())
    
    for paragraph in template_document.paragraphs:
        if '{{' in paragraph.text:
            for variable_key, variable_value in variables.items():
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

    for table in template_document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{' in paragraph.text:
                        for variable_key, variable_value in variables.items():
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

    if template_file_path in ['00 Подписанный Дог серт Исходник.docx',
                              '01 Дог серт Исходник.docx',
                              '03 Договор ИК Исходник.docx',
                              '01 Заявка Исходник.docx']:
        if variables['worksheets_count'] > 1:
            for table in template_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        currentSheet = theFile.worksheets[0]
                        # Это условие для добавления строк таблицы в договоре и ИК
                        if '{{contr_tabl}}' in cell.text:
                            cell.text = str(currentSheet['B18'].value)
                            for work in range(1, variables['worksheets_count']):
                                currentSheet = theFile.worksheets[work]
                                new_row = table.add_row()
                                new_cells = new_row.cells
                                new_cells[0].text = str(work + 1)
                                # Эта строка центрирует текст
                                # Обращаться надо именно через .paragraphs[0].alignment
                                new_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                new_cells[1].text = str(currentSheet['B16'].value)
                                new_cells[2].text = str(currentSheet['B18'].value)

                        # Это условие для добавления строк таблицы в заявке
                        if '{{request_tabl}}' in cell.text:
                            cell.text = str(currentSheet['B21'].value) + '\n' + str(currentSheet['B18'].value)
                            for work in range(1, variables['worksheets_count']):
                                currentSheet = theFile.worksheets[work]
                                new_row = table.add_row()
                                new_cells = new_row.cells
                                new_cells[0].text = str(work + 1)
                                # Эта строка центрирует текст
                                # Обращаться надо именно через .paragraphs[0].alignment
                                new_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                new_cells[1].text = str(currentSheet['B16'].value)
                                new_cells[2].text = str(currentSheet['B21'].value) + '\n' + str(currentSheet['B18'].value)
        # Здесь убираются {{contr_tabl}} и {{request_tabl}}
        else:
            for table in template_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        currentSheet = theFile.worksheets[0]
                        # Это условие для добавления строк таблицы в договоре и ИК
                        if '{{contr_tabl}}' in cell.text:
                            cell.text = str(currentSheet['B18'].value)

                        # Это условие для добавления строк таблицы в заявке
                        if '{{request_tabl}}' in cell.text:
                            cell.text = str(currentSheet['B21'].value) + '\n' + str(currentSheet['B18'].value)


    # if variables['worksheets_count'] > 1:
    #     for table in template_document.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 if '{{contr_tabl}}' in cell.text:
    #                     # Add a new row to the table after the row containing the marker


    template_document.save(output_file_path)
    print('OK!!!!!!!!!!!!!!!!!!!', output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        text_runs = [(run.text, run.font.name, run.font.size, run.bold, run.italic, run.underline) for run in paragraph.runs]
        full_text = ''.join([text for text, _, _, _, _, _ in text_runs])
        new_text = full_text.replace(key, value)

        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''

        # Create a new run to hold the entire new text
        new_run = paragraph.add_run(new_text)

        # Apply formatting from the first run
        if text_runs:
            _, font_name, font_size, bold, italic, underline = text_runs[0]
            new_run.font.name = font_name or 'Times New Roman'
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline


def reference_data_xlsx(input_reference, output_reference, variables):
    theFile = openpyxl.load_workbook(input_reference)
    currentSheet = theFile['На регистрацию сертиф.']
    currentSheet['D8'].value = f"{variables['{{WORK_NUMBER}}']} от {variables['{{CONTR_DATE}}']}.{variables['{{CONTR_MONTH}}']}.{variables['{{CONTR_YEAR}}']}г."
    currentSheet['D11'].value = f"{variables['{{EXPERT_LASTNAME}}']} {variables['{{EXPERT_FIRSTNAME_SHORT}}']}{variables['{{EXPERT_SECNAME_SHORT}}']}"
    currentSheet['D12'].value = f"RU.МСС.Э.{variables['{{EXPERT_REG_NUMBER}}']}"
    currentSheet['D13'].value = variables['{{IL_NAME}}']
    currentSheet['D14'].value = f"RU.МСС.Л.{variables['{{IL_REG_NUMBER}}']}"
    currentSheet['D15'].value = variables['{{BUSINESS_FORM}}']
    currentSheet['D16'].value = variables['{{COMPANY_NAME}}']
    currentSheet['D17'].value = variables['{{JUR_ADDRESS}}']
    currentSheet['D18'].value = variables['{{TEL}}']
    currentSheet['D19'].value = variables['{{E-MAIL}}']
    currentSheet['D20'].value = f"{variables['{{DIR_LASTNAME}}']} {variables['{{DIR_FIRSTNAME}}']} {variables['{{DIR_SECNAME}}']}"
    currentSheet['D21'].value = variables['{{INN}}']
    currentSheet['D22'].value = variables['{{KPP}}']
    currentSheet['D27'].value = variables['{{CERT_NAME}}']
    currentSheet['D28'].value = variables['{{OKPD}}']
    currentSheet['D31'].value = variables['{{STANDART_SHORT}}']
    currentSheet['D32'].value = variables['{{STANDART_FULL}}']
    currentSheet['D33'].value = variables['{{WORK_NUMBER}}']
    currentSheet['D34'].value = f"{variables['{{ISSUE_DECISION_DATE}}']}.{variables['{{ISSUE_DECISION_MONTH}}']}.{variables['{{ISSUE_DECISION_YEAR}}']}"
    currentSheet['D35'].value = f"{variables['{{CERTIFICARTE_START_DATE}}']}.{variables['{{CERTIFICARTE_START_MONTH}}']}.{variables['{{CERTIFICARTE_START_YEAR}}']}"
    currentSheet['D36'].value = f"{variables['{{CERTIFICARTE_DURATION}}']} {'года' if int(variables['{{CERTIFICARTE_DURATION}}']) <= 4 else 'лет'}"
    currentSheet['D37'].value = f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 1)}"
    currentSheet['D38'].value = f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 2)}"
    currentSheet['D39'].value = "" if int(variables['{{CERTIFICARTE_DURATION}}']) <= 3 else f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 3)}"
    currentSheet['D40'].value = "" if int(variables['{{CERTIFICARTE_DURATION}}']) <= 4 else f"15.{variables['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables['{{CERTIFICARTE_START_YEAR}}']) + 4)}"
    currentSheet['D42'].value = f"{variables['{{SAMPLE_ACT_DATE}}']}.{variables['{{SAMPLE_ACT_MONTH}}']}.{variables['{{SAMPLE_ACT_YEAR}}']}"
    currentSheet['D44'].value = variables['{{PROTOCOL_NUMBER}}']
    currentSheet['D45'].value = f"{variables['{{PROTOCOL_DATE}}']}.{variables['{{PROTOCOL_MONTH}}']}.{variables['{{PROTOCOL_YEAR}}']}"
    currentSheet['D46'].value = f"{variables['{{TESTER_NAME}}']}"
    currentSheet['D47'].value = f"{variables['{{PROD_ANALYSE_DATE}}']}.{variables['{{PROD_ANALYSE_MONTH}}']}.{variables['{{PROD_ANALYSE_YEAR}}']} № {variables['{{WORK_NUMBER}}']}"
    currentSheet['D48'].value = f"{variables['{{CONTRACT_NUMBER}}']}-{variables['{{CONTRACT_YEAR}}']}"
    currentSheet['D49'].value = f"{variables['{{CONTR_DATE}}']}.{variables['{{CONTR_MONTH}}']}.{variables['{{CONTR_YEAR}}']}г."
    currentSheet['D50'].value = f"{variables['{{CONTRACT_SUM}}']}"
    currentSheet['D51'].value = f"{variables['{{CONTRACT_OS_FULL_SUM}}']}"
    currentSheet['D52'].value = f"{variables['{{CONTRACT_OS_SUM}}']}"
    currentSheet['D53'].value = f"{variables['{{CONTRACT_IL_FULL_SUM}}']}"
    currentSheet['D54'].value = f"{variables['{{CONTRACT_IL_SUM}}']}"
    theFile.save(output_reference)


def ik_output(input_reference, excel_file_path, company, variables):
    """Создаёт excel заготовку для инспекции"""

    IK_file = openpyxl.load_workbook(input_reference)
    for work in range(variables['worksheets_count']):
        variables = load_variables_from_sheet(excel_file_path, sheet_number=work)
        currentSheet = IK_file.worksheets[work]
        currentSheet.title = f'Инсп 0{work+1}'
        currentSheet['B2'].value = '1'
        currentSheet['B3'].value = variables['{{CONTRACT_NUMBER}}']
        currentSheet['B4'].value = variables['{{CONTRACT_YEAR}}']
        currentSheet['B5'].value = f"{variables['{{CONTRACT_NUMBER}}']}-{work+1}"
        currentSheet['B6'].value = variables['{{IK&CERT_DAY}}']
        currentSheet['B7'].value = variables['{{IK&CERT_MONTH}}']
        currentSheet['B8'].value = variables['{{IK&CERT_YEAR}}']
        currentSheet['B9'].value = variables['{{BUSINESS_FORM}}']
        currentSheet['B10'].value = variables['{{COMPANY_NAME}}']
        currentSheet['B11'].value = variables['{{DIR_LASTNAME}}']  # Фамилия
        currentSheet['B12'].value = variables['{{DIR_FIRSTNAME}}']  # Имя
        currentSheet['B13'].value = variables['{{DIR_SECNAME}}']  # Отчество
        currentSheet['B14'].value = ""
        currentSheet['B15'].value = variables['{{IK&CERT_DAY}}']
        currentSheet['B16'].value = variables['{{IK&CERT_MONTH}}']
        currentSheet['B17'].value = variables['{{IK&CERT_YEAR}}']
        currentSheet['B18'].value = variables['{{CERT_NAME}}']
        currentSheet['B19'].value = variables['{{OKPD}}']
        currentSheet['B20'].value = variables['{{STANDART_FULL}}']
        currentSheet['B21'].value = variables['{{CONTRACT_IK_SUM}}']
        currentSheet['B22'].value = variables['{{CONTRACT_IK_SUM}}']
        currentSheet['B23'].value = ""
        currentSheet['B24'].value = variables['{{INN}}']
        currentSheet['B25'].value = variables['{{KPP}}']
        currentSheet['B26'].value = variables['{{JUR_ADDRESS}}']
        currentSheet['B27'].value = variables['{{RAS_ACCOUNT}}']
        currentSheet['B28'].value = variables['{{BANK_NAME}}']
        currentSheet['B29'].value = variables['{{KORR_ACCOUNT}}']
        currentSheet['B30'].value = variables['{{BIK}}']
        currentSheet['B31'].value = variables['{{OGRN}}']
        currentSheet['B32'].value = variables['{{TEL}}']
        currentSheet['B33'].value = ""
        currentSheet['B34'].value = ""
        currentSheet['B35'].value = ""
        currentSheet['B36'].value = ""
        currentSheet['B37'].value = ""
        currentSheet['B38'].value = ""
        IK_file.save(f"{company}/Инспекция Исходник {variables['{{COMPANY_NAME}}']}.xlsx")
    print('OK!!!!!!!!!!!!!!!!')


if __name__ == '__main__':
    excel_file_path = 'Сертификация много исходник.xlsx'
    currentSheet = theFile.worksheets[0]
    variables = load_variables_from_sheet(excel_file_path, sheet_number=0)
    company = variables['{{COMPANY_NAME}}']
    BASE_DIR = Path(__file__).parent
    main_dir = BASE_DIR / f'{company}'
    main_dir.mkdir(exist_ok=True)
    main_dir = BASE_DIR / f'{company}' / 'Договора'
    main_dir.mkdir(exist_ok=True)
    main_dir = BASE_DIR / f'{company}' / 'Комплект документов'
    main_dir.mkdir(exist_ok=True)
    works = len(theFile.worksheets)
    if works > 1:
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '01 Заявка'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '02 Распоряжения'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '03 Испытания'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '04 Акты обследования производства'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '05 Заключения'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '06 Решения о выдаче'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '07 Макеты сертификатов'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / f'{company}' / 'Комплект документов' / '08 Справки'
        main_dir.mkdir(exist_ok=True)
    main('00 Подписанный Дог серт Исходник.docx', f"{company}/Договора/00 Подписанный Дог серт {company}.docx", variables)
    main('01 Дог серт Исходник.docx', f"{company}/Договора/01 Дог серт {company}.docx", variables)
    main('02 Счет Исходник.docx', f"{company}/Договора/02 Счет {company}.docx", variables)
    main('03 Договор ИК Исходник.docx', f"{company}/Договора/03 Договор ИК {variables['{{CONTRACT_NUMBER}}']}-1-{variables['{{CONTRACT_YEAR}}']} {company}.docx", variables)
    main('04 АКТ приемки Исходник.docx', f"{company}/Договора/04 АКТ приемки ХХ {company}.docx", variables)
    main('01 Заявка Исходник.docx', f"{company}/Комплект документов{'/01 Заявка' if works > 1 else ''}/01 Заявка {company}.docx", variables)
    # convert(f"{company}/Договора/00 Подписанный Дог серт {company}.docx")
    # convert(f"{company}/Договора/02 Счет {company}.docx")
    works = len(theFile.worksheets)
    for work in range(works):
        sheet_number = work
        variables = load_variables_from_sheet(excel_file_path, sheet_number)
        # add - добавка -1, -2 и т.д. к номеру работы, если больше одной работы
        add = ('' if works <= 1 else '-' + str(work + 1))
        main('02 Распоряжение Исходник.docx', f"{company}/Комплект документов{'/02 Распоряжения' if works > 1 else ''}/02{add} Распоряжение {company}.docx", variables)
        main('03 Акт отбора Исходник.docx', f"{company}/Комплект документов{'/03 Испытания' if works > 1 else ''}/03{add} Акт отбора {company}.docx", variables)
        main('03 Протокол Исходник.docx', f"{company}/Комплект документов{'/03 Испытания' if works > 1 else ''}/03{add} Протокол {company}.docx", variables)
        main('04 Анализ производства Исходник.docx', f"{company}/Комплект документов{'/04 Акты обследования производства' if works > 1 else ''}/04{add} Анализ производства {company}.docx", variables)
        main('05 Заключение Исходник.docx', f"{company}/Комплект документов{'/05 Заключения' if works > 1 else ''}/05{add} Заключение {company}.docx", variables)
        main('06 Решение о выдаче Исходник.docx', f"{company}/Комплект документов{'/06 Решения о выдаче' if works > 1 else ''}/06{add} Решение о выдаче {company}.docx", variables)
        main('07 Макет сертификата Исходник.docx', f"{company}/Комплект документов{'/07 Макеты сертификатов' if works > 1 else ''}/07{add} Макет сертификата {company}.docx", variables)
        reference_data_xlsx('08 Справка Исходник.xlsx', f"{company}/Комплект документов{'/08 Справки' if works > 1 else ''}/08{add} Справка {company}.xlsx", variables)
    ik_output('Инспекция Заготовка.xlsx', excel_file_path, company, variables)

