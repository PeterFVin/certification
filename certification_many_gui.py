from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx2pdf import convert
import openpyxl


MONTHS = {
    '01': 'январь',
    '02': 'февраль',
    '03': 'март',
    '04': 'апрель',
    '05': 'май',
    '06': 'июнь',
    '07': 'июль',
    '08': 'август',
    '09': 'сентябрь',
    '10': 'октябрь',
    '11': 'ноябрь',
    '12': 'декабрь',
}

counter1 = 0

def main(template_file_path, output_file_path, variables, counter1):
    template_document = Document(template_file_path)

    print('!!!!!!!!!!!!!!!!!!!!!!!!!!')
    for paragraph in template_document.paragraphs:
        if '{{' in paragraph.text:
            for variable_key, variable_value in variables[counter1].items():
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

    for table in template_document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{' in paragraph.text:
                        for variable_key, variable_value in variables[counter1].items():
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

    if template_file_path in ['00 Подписанный Дог серт Исходник.docx',
                              '01 Дог серт Исходник.docx',
                              '03 Договор ИК Исходник.docx',
                              '01 Заявка Исходник.docx']:
        if variables[counter1]['worksheets_count'] > 1:
            for table in template_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # currentSheet = theFile.worksheets[0]
                        # Это условие для добавления строк таблицы в договоре и ИК
                        if '{{contr_tabl}}' in cell.text:
                            cell.text = variables['OKPD'][0]
                            for work in range(1, variables[counter1]['worksheets_count']):
                                # currentSheet = theFile.worksheets[work]
                                new_row = table.add_row()
                                new_cells = new_row.cells
                                new_cells[0].text = str(work + 1)
                                # Эта строка центрирует текст
                                # Обращаться надо именно через .paragraphs[0].alignment
                                new_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                new_cells[1].text = variables['CERT_NAME'][work]
                                new_cells[2].text = variables['OKPD'][work]

                        # Это условие для добавления строк таблицы в заявке
                        if '{{request_tabl}}' in cell.text:
                            cell.text = variables['STANDART_FULL'][0] + '\n' + variables['OKPD'][0]
                            for work in range(1, variables[counter1]['worksheets_count']):
                                # currentSheet = theFile.worksheets[work]
                                new_row = table.add_row()
                                new_cells = new_row.cells
                                new_cells[0].text = str(work + 1)
                                # Эта строка центрирует текст
                                # Обращаться надо именно через .paragraphs[0].alignment
                                new_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                new_cells[1].text = variables['CERT_NAME'][work]
                                new_cells[2].text = variables['STANDART_FULL'][work] + '\n' + variables['OKPD'][work]
        # Здесь убираются {{contr_tabl}} и {{request_tabl}}
        else:
            for table in template_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # currentSheet = theFile.worksheets[0]
                        # Это условие для добавления строк таблицы в договоре и ИК
                        if '{{contr_tabl}}' in cell.text:
                            cell.text = variables['OKPD'][0]

                        # Это условие для добавления строк таблицы в заявке
                        if '{{request_tabl}}' in cell.text:
                            cell.text = variables['STANDART_FULL'][0] + '\n' + variables['OKPD'][0]

    template_document.save(output_file_path)
    print('OK!!!!!!!!!!!!!!!!!!!', output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        text_runs = [(run.text, run.font.name, run.font.size, run.bold, run.italic, run.underline) for run in paragraph.runs]
        full_text = ''.join([text for text, _, _, _, _, _ in text_runs])
        new_text = full_text.replace(key, value)

        for run in paragraph.runs:
            run.text = ''

        new_run = paragraph.add_run(new_text)

        if text_runs:
            _, font_name, font_size, bold, italic, underline = text_runs[0]
            new_run.font.name = font_name or 'Times New Roman'
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline


def reference_data_xlsx(input_reference, output_reference, variables, counter1):
    theFile = openpyxl.load_workbook(input_reference)
    currentSheet = theFile['На регистрацию сертиф.']
    currentSheet['D8'].value = f"{variables[counter1]['{{WORK_NUMBER}}']} от {variables[counter1]['{{CONTR_DATE}}']}.{variables[counter1]['{{CONTR_MONTH}}']}.{variables[counter1]['{{CONTR_YEAR}}']}г."
    currentSheet['D11'].value = f"{variables[counter1]['{{EXPERT_LASTNAME}}']} {variables[counter1]['{{EXPERT_FIRSTNAME_SHORT}}']}{variables[counter1]['{{EXPERT_SECNAME_SHORT}}']}"
    currentSheet['D12'].value = f"RU.МСС.Э.{variables[counter1]['{{EXPERT_REG_NUMBER}}']}"
    currentSheet['D13'].value = variables[counter1]['{{IL_NAME}}']
    currentSheet['D14'].value = f"RU.МСС.Л.{variables[counter1]['{{IL_REG_NUMBER}}']}"
    currentSheet['D15'].value = variables[counter1]['{{BUSINESS_FORM}}']
    currentSheet['D16'].value = variables[counter1]['{{COMPANY_NAME}}']
    currentSheet['D17'].value = variables[counter1]['{{JUR_ADDRESS}}']
    currentSheet['D18'].value = variables[counter1]['{{TEL}}']
    currentSheet['D19'].value = variables[counter1]['{{E-MAIL}}']
    currentSheet['D20'].value = f"{variables[counter1]['{{DIR_LASTNAME}}']} {variables[counter1]['{{DIR_FIRSTNAME}}']} {variables[counter1]['{{DIR_SECNAME}}']}"
    currentSheet['D21'].value = variables[counter1]['{{INN}}']
    currentSheet['D22'].value = variables[counter1]['{{KPP}}']
    currentSheet['D27'].value = variables[counter1]['{{CERT_NAME}}']
    currentSheet['D28'].value = variables[counter1]['{{OKPD}}']
    currentSheet['D31'].value = variables[counter1]['{{STANDART_SHORT}}']
    currentSheet['D32'].value = variables[counter1]['{{STANDART_FULL}}']
    currentSheet['D33'].value = variables[counter1]['{{WORK_NUMBER}}']
    currentSheet['D34'].value = f"{variables[counter1]['{{ISSUE_DECISION_DATE}}']}.{variables[counter1]['{{ISSUE_DECISION_MONTH}}']}.{variables[counter1]['{{ISSUE_DECISION_YEAR}}']}"
    currentSheet['D35'].value = f"{variables[counter1]['{{CERTIFICARTE_START_DATE}}']}.{variables[counter1]['{{CERTIFICARTE_START_MONTH}}']}.{variables[counter1]['{{CERTIFICARTE_START_YEAR}}']}"
    currentSheet['D36'].value = f"{variables[counter1]['{{CERTIFICARTE_DURATION}}']} {'года' if int(variables[counter1]['{{CERTIFICARTE_DURATION}}']) <= 4 else 'лет'}"
    currentSheet['D37'].value = f"15.{variables[counter1]['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables[counter1]['{{CERTIFICARTE_START_YEAR}}']) + 1)}"
    currentSheet['D38'].value = f"15.{variables[counter1]['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables[counter1]['{{CERTIFICARTE_START_YEAR}}']) + 2)}"
    currentSheet['D39'].value = "" if int(variables[counter1]['{{CERTIFICARTE_DURATION}}']) <= 3 else f"15.{variables[counter1]['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables[counter1]['{{CERTIFICARTE_START_YEAR}}']) + 3)}"
    currentSheet['D40'].value = "" if int(variables[counter1]['{{CERTIFICARTE_DURATION}}']) <= 4 else f"15.{variables[counter1]['{{CERTIFICARTE_START_MONTH}}']}.{str(int(variables[counter1]['{{CERTIFICARTE_START_YEAR}}']) + 4)}"
    currentSheet['D42'].value = f"{variables[counter1]['{{SAMPLE_ACT_DATE}}']}.{variables[counter1]['{{SAMPLE_ACT_MONTH}}']}.{variables[counter1]['{{SAMPLE_ACT_YEAR}}']}"
    currentSheet['D44'].value = variables[counter1]['{{PROTOCOL_NUMBER}}']
    currentSheet['D45'].value = f"{variables[counter1]['{{PROTOCOL_DATE}}']}.{variables[counter1]['{{PROTOCOL_MONTH}}']}.{variables[counter1]['{{PROTOCOL_YEAR}}']}"
    currentSheet['D46'].value = f"{variables[counter1]['{{TESTER_NAME}}']}"
    currentSheet['D47'].value = f"{variables[counter1]['{{PROD_ANALYSE_DATE}}']}.{variables[counter1]['{{PROD_ANALYSE_MONTH}}']}.{variables[counter1]['{{PROD_ANALYSE_YEAR}}']} № {variables[counter1]['{{WORK_NUMBER}}']}"
    currentSheet['D48'].value = f"{variables[counter1]['{{CONTRACT_NUMBER}}']}-{variables[counter1]['{{CONTRACT_YEAR}}']}"
    currentSheet['D49'].value = f"{variables[counter1]['{{CONTR_DATE}}']}.{variables[counter1]['{{CONTR_MONTH}}']}.{variables[counter1]['{{CONTR_YEAR}}']}г."
    currentSheet['D50'].value = f"{variables[counter1]['{{CONTRACT_SUM}}']}"
    currentSheet['D51'].value = f"{variables[counter1]['{{CONTRACT_OS_FULL_SUM}}']}"
    currentSheet['D52'].value = f"{variables[counter1]['{{CONTRACT_OS_SUM}}']}"
    currentSheet['D53'].value = f"{variables[counter1]['{{CONTRACT_IL_FULL_SUM}}']}"
    currentSheet['D54'].value = f"{variables[counter1]['{{CONTRACT_IL_SUM}}']}"
    theFile.save(output_reference)

def main_func(variables, documents_single, documents_many, spravka, counter1=counter1):
    company = variables[1]['{{COMPANY_NAME}}']
    contr_num = variables[1]['{{CONTRACT_NUMBER}}']
    year_num = variables[1]['{{CONTRACT_YEAR}}']
    BASE_DIR = Path(__file__).parent
    direct = f'{contr_num}-{year_num} {company}'
    main_dir = BASE_DIR / direct
    main_dir.mkdir(exist_ok=True)
    main_dir = BASE_DIR / direct / 'Договора'
    main_dir.mkdir(exist_ok=True)
    main_dir = BASE_DIR / direct / 'Комплект документов'
    main_dir.mkdir(exist_ok=True)
    works = variables[1]['worksheets_count']
    if works > 1:
        main_dir = BASE_DIR / direct / 'Комплект документов' / '01 Заявка'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '02 Распоряжения'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '03 Испытания'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '04 Акты обследования производства'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '05 Заключения'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '06 Решения о выдаче'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '07 Макеты сертификатов'
        main_dir.mkdir(exist_ok=True)
        main_dir = BASE_DIR / direct / 'Комплект документов' / '08 Справки'
        main_dir.mkdir(exist_ok=True)
    documents_one = {
        0: ('00 Подписанный Дог серт Исходник.docx', f"{direct}/Договора/00 Подписанный Дог серт {company}.docx"),
        1: ('01 Дог серт Исходник.docx', f"{direct}/Договора/01 Дог серт {contr_num}-{year_num} {company}.docx"),
        2: ('02 Счет Исходник.docx', f"{direct}/Договора/02 Счет {variables[1]['{{BILL_NUMBER}}']} от {variables[1]['{{CONTR_DATE}}']}.{variables[1]['{{CONTR_MONTH}}']}.{variables[1]['{{CONTR_YEAR}}']} {company}.docx"),
        3: ('03 Договор ИК Исходник.docx', f"{direct}/Договора/03 Договор ИК {contr_num}-1-{year_num} {company}.docx"),
        4: ('04 АКТ приемки Исходник.docx', f"{direct}/Договора/04 АКТ приемки ХХ {company}.docx"),
        5: ('01 Заявка Исходник.docx', f"{direct}/Комплект документов{'/01 Заявка' if works > 1 else ''}/01 Заявка {company}.docx")
    }
    for key, value in documents_single.items():
        if value:
            main(documents_one[key][0], documents_one[key][1], variables, counter1=1)
    for work in range(works):
        counter1 += 1
        sheet_number = work
        # add - добавка -1, -2 и т.д. к номеру работы, если больше одной работы
        add = ('' if works <= 1 else '-' + str(work + 1))
        documents_multi = {
            0: ('02 Распоряжение Исходник.docx', f"{direct}/Комплект документов{'/02 Распоряжения' if works > 1 else ''}/02{add} Распоряжение {company}.docx"),
            1: ('03 Акт отбора Исходник.docx', f"{direct}/Комплект документов{'/03 Испытания' if works > 1 else ''}/03{add} Акт отбора {company}.docx"),
            2: ('03 Протокол Исходник.docx', f"{direct}/Комплект документов{'/03 Испытания' if works > 1 else ''}/03{add} Протокол {company}.docx"),
            3: ('04 Анализ производства Исходник.docx', f"{direct}/Комплект документов{'/04 Акты обследования производства' if works > 1 else ''}/04{add} Анализ производства {company}.docx"),
            4: ('05 Заключение Исходник.docx', f"{direct}/Комплект документов{'/05 Заключения' if works > 1 else ''}/05{add} Заключение {company}.docx"),
            5: ('06 Решение о выдаче Исходник.docx', f"{direct}/Комплект документов{'/06 Решения о выдаче' if works > 1 else ''}/06{add} Решение о выдаче {company}.docx"),
            6: ('07 Макет сертификата Исходник.docx', f"{direct}/Комплект документов{'/07 Макеты сертификатов' if works > 1 else ''}/07{add} Макет сертификата {company}.docx")
        }
        
        for key, value in documents_many.items():
            if value:
                main(documents_multi[key][0], documents_multi[key][1], variables, counter1)
            if spravka:
                reference_data_xlsx('08 Справка Исходник.xlsx', f"{direct}/Комплект документов{'/08 Справки' if works > 1 else ''}/08{add} Справка {company}.xlsx", variables, counter1)
